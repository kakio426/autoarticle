import streamlit as st
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
import io
import os
import requests
import datetime
import pandas as pd
import json
import uuid
from PIL import Image, ImageDraw, ImageFont
import zipfile
import re

# ==========================================
# 0. 초기 설정 및 디렉토리 관리
# ==========================================
st.set_page_config(layout="wide", page_title="AutoSchoolArticle: 학교 소식지 자동화")

ARCHIVE_DIR = "archive"
IMAGE_DIR = os.path.join(ARCHIVE_DIR, "images")
DATA_FILE = os.path.join(ARCHIVE_DIR, "data.csv")

for folder in [ARCHIVE_DIR, IMAGE_DIR]:
    if not os.path.exists(folder):
        os.makedirs(folder, exist_ok=True)

if not os.path.exists(DATA_FILE):
    df_empty = pd.DataFrame(columns=["id", "date", "school", "grade", "event_name", "location", "tone", "keywords", "title", "content", "images", "hashtags"])
    df_empty.to_csv(DATA_FILE, index=False, encoding='utf-8-sig')

# 한글 폰트 설정
FONT_URL = "https://github.com/google/fonts/raw/main/ofl/nanumgothic/NanumGothic-Regular.ttf"
FONT_PATH = "NanumGothic-Regular.ttf"

def ensure_font():
    if not os.path.exists(FONT_PATH):
        try:
            response = requests.get(FONT_URL)
            with open(FONT_PATH, "wb") as f:
                f.write(response.content)
        except:
            pass

ensure_font()

# 디자인 테마 (컬러칩 미리보기용 데이터 포함)
THEMES = {
    "웜 & 플레이풀": {"main": (255, 140, 66), "sub": (255, 251, 240), "accent": (6, 214, 160), "hex": "#FF8C42"},
    "꿈꾸는 파랑": {"main": (0, 80, 150), "sub": (230, 240, 255), "accent": (0, 120, 215), "hex": "#005096"},
    "발랄한 노랑": {"main": (255, 180, 0), "sub": (255, 250, 230), "accent": (255, 140, 0), "hex": "#FFB400"},
    "산뜻한 민트": {"main": (0, 168, 107), "sub": (235, 250, 245), "accent": (0, 128, 90), "hex": "#00A86B"}
}

def apply_custom_style():
    st.markdown("""
        <style>
        @import url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_two@1.0/NanumSquareRound.woff');
        
        :root {
            --primary: #FF8C42;
            --bg: #FFFBF0;
            --text: #4A4A4A;
            --radius: 16px;
            --input-radius: 12px;
        }

        /* 1. 기본 배경 */
        .stApp { background-color: var(--bg) !important; }

        /* 2. 텍스트 요소 통합 관리 */
        h1, h2, h3, h4, h5, h6, p, label, .stMarkdown, .stText, .stCaption, 
        .stButton button, .stCheckbox label span, [data-baseweb="select"] * {
            font-family: 'NanumSquareRound', sans-serif !important;
            color: var(--text) !important;
        }

        /* 3. 입력창 찌꺼기 및 '옹졸한 박스' 제거 (근본 해결) */
        /* [data-baseweb] 하위의 모든 입력 필드 잔상 제거 */
        [data-baseweb="input"] input, [data-baseweb="textarea"] textarea, [data-baseweb="select"] input {
            background-color: transparent !important;
            border: none !important;
            box-shadow: none !important;
            outline: none !important;
            padding: 0 !important;
        }

        /* 4. 실제 보이는 박스 디자인 (깔끔한 unit) */
        /* 일반 입력창 및 날짜 입력창: 내부 여백 필요 */
        .stTextInput > div > div, 
        .stTextArea > div > div, 
        .stDateInput > div > div {
            background-color: white !important;
            border: 1.5px solid #FFE0B3 !important;
            border-radius: var(--input-radius) !important;
            padding: 4px 12px !important;
        }
        
        /* 셀렉트박스 (참여 학년 등): 외부 패딩 제거 (잘림 방지) */
        .stSelectbox > div > div {
            background-color: white !important;
            border: 1.5px solid #FFE0B3 !important;
            border-radius: var(--input-radius) !important;
            padding: 0 !important; /* ★ 중요: 외부 패딩 제거 ★ */
        }

        /* 셀렉트박스 내부 텍스트 정렬 및 높이 확보 */
        div[data-baseweb="select"] {
            min-height: 48px !important;
        }
        div[data-baseweb="select"] > div {
            padding: 0 12px !important;
            min-height: 48px !important;
            display: flex !important;
            align-items: center !important; /* 수직 중앙 정렬 */
        }

        /* 5. 포커스 시 테두리 강조 */
        .stTextInput > div > div:focus-within, 
        .stTextArea > div > div:focus-within, 
        .stSelectbox > div > div:focus-within,
        .stDateInput > div > div:focus-within {
            border-color: var(--primary) !important;
            border-width: 2px !important;
        }

        /* 6. 아이콘 보호 및 헤더 클리닝 */
        [data-testid="stIcon"], i, svg { font-family: inherit !important; }
        header[data-testid="stHeader"] { background-color: transparent !important; }
        header[data-testid="stHeader"] * { color: transparent !important; }
        header[data-testid="stHeader"] svg { fill: var(--primary) !important; }

        /* 7. 몽글몽글 버튼 */
        .stButton>button {
            border-radius: var(--radius) !important;
            background-color: var(--primary) !important;
            color: white !important;
            border: none !important;
            box-shadow: 0 4px 0px #E67E30 !important;
            padding: 0.6rem 2rem !important;
            font-weight: bold !important;
        }
        .stButton>button:hover {
            transform: translateY(-2px) !important;
            box-shadow: 0 6px 0px #E67E30 !important;
        }

        /* 8. 카드 및 사이드바 (색상 복구 중요) */
        div[data-testid="stVerticalBlock"] > div[style*="border"] {
            background-color: white !important;
            border: 1px solid #FFEBB3 !important;
            border-radius: var(--radius) !important;
            box-shadow: 0 8px 30px rgba(230, 126, 34, 0.03) !important;
        }

        /* 푸른색을 지우고 웜 베이지로 복구 */
        [data-testid="stSidebar"] {
            background-color: #FFF9E6 !important;
            border-right: 2px solid #FFEBB3 !important;
        }

        /* 9. [NEW] 전문 기사 포맷 전용 스타일 (Compact) */
        .article-card {
            background: white;
            border-radius: 20px;
            padding: 35px;
            border: 1px solid #FFE0B3;
            box-shadow: 0 10px 40px rgba(0,0,0,0.02);
            max-width: 850px; /* 너무 퍼지지 않게 제한 */
            margin: 0 auto;
        }
        .article-header {
            margin-bottom: 25px;
            border-bottom: 2px solid #FFF5E6;
            padding-bottom: 20px;
        }
        .article-title {
            font-size: 1.8rem !important; /* 살짝 축소 */
            font-weight: 800 !important;
            color: #E67E30 !important;
            margin-bottom: 12px !important;
            line-height: 1.3 !important;
        }
        .badge-container {
            display: flex;
            gap: 10px;
            flex-wrap: wrap;
        }
        .badge {
            background: #FFF2E6;
            color: #FF8C42;
            padding: 4px 12px;
            border-radius: 50px;
            font-size: 0.8rem;
            font-weight: 600;
            display: flex;
            align-items: center;
            gap: 5px;
            border: 1px solid #FFE4CC;
        }
        .article-content {
            font-size: 1.05rem !important;
            line-height: 1.7 !important;
            color: #4A4A4A !important;
            white-space: pre-wrap;
        }
        /* 이미지 그리드 시스템 (크기 최적화) */
        .img-grid {
            display: grid;
            gap: 12px;
            margin: 25px 0;
            max-width: 750px;
            margin-left: auto;
            margin-right: auto;
        }
        .img-item {
            width: 100%;
            height: 220px; /* 이미지 높이 축소 */
            object-fit: cover;
            border-radius: 12px;
            border: 1px solid #FFE0B3;
        }
        .grid-1 { grid-template-columns: 1fr; }
        .grid-1 .img-item { height: 320px; } /* 1장일 때도 너무 크지 않게 */
        .grid-2 { grid-template-columns: 1fr 1fr; }
        .grid-3 { grid-template-areas: "main main" "sub1 sub2"; }
        .grid-3 .img-item:first-child { grid-area: main; height: 300px; }
        .grid-4 { grid-template-columns: 1fr 1fr; }
        </style>
    """, unsafe_allow_html=True)

def show_stepper(step_idx):
    steps = ["1. 정보 입력", "2. AI 초안 생성", "3. 편집 및 보존", "4. 뉴스레터 발행"]
    cols = st.columns(len(steps))
    for i, step in enumerate(steps):
        is_active = i == step_idx
        color = "#FF8C42" if is_active else "#EBD4BD"
        bg = "#FFF2E6" if is_active else "transparent"
        cols[i].markdown(f"""
            <div style="text-align:center; padding:10px; border-radius:16px; background:{bg}; border: 2px dashed {color if is_active else 'transparent'}">
                <span style="color:{color}; font-weight:{'bold' if is_active else 'normal'}">{step}</span>
            </div>
        """, unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)

# ==========================================
# 1. 로직: 아카이브 관리
# ==========================================
def save_to_archive(data, uploaded_files):
    df = pd.read_csv(DATA_FILE, encoding='utf-8-sig')
    
    # 이미지 저장
    image_paths = []
    if uploaded_files:
        for file in uploaded_files:
            img_id = str(uuid.uuid4())
            # Handle both Streamlit UploadedFile and potential local paths
            ext = file.name.split('.')[-1] if hasattr(file, 'name') else 'png'
            file_path = os.path.join(IMAGE_DIR, f"{img_id}.{ext}")
            with open(file_path, "wb") as f:
                f.write(file.getbuffer() if hasattr(file, 'getbuffer') else file)
            image_paths.append(file_path)
    
    new_data = {
        "id": str(uuid.uuid4()),
        "date": data['date'],
        "school": data['school'],
        "grade": data['grade'],
        "event_name": data['event_name'],
        "location": data['location'],
        "tone": data['tone'],
        "keywords": data['keywords'],
        "title": data['title'],
        "content": data['content'],
        "images": json.dumps(image_paths),
        "hashtags": json.dumps(data.get('hashtags', []))
    }
    
    df = pd.concat([df, pd.DataFrame([new_data])], ignore_index=True)
    df.to_csv(DATA_FILE, index=False, encoding='utf-8-sig')
    return new_data

def update_archive(target_id, updated_title, updated_content):
    df = pd.read_csv(DATA_FILE, encoding='utf-8-sig')
    idx = df[df['id'] == target_id].index
    if not idx.empty:
        df.loc[idx, 'title'] = updated_title
        df.loc[idx, 'content'] = updated_content
        df.to_csv(DATA_FILE, index=False, encoding='utf-8-sig')
        return True
    return False

# ==========================================
# 2. 로직: AI 기사 작성 (Gemini)
# ==========================================
def generate_article_gemini(api_key, topic_data):
    if not api_key:
        return "API 키가 필요합니다.", "사이드바에 API 키를 입력해주세요."
    
    try:
        genai.configure(api_key=api_key)
        # 최신 고성능 Flash 모델 설정 (Gemini 3 Flash Preview - 2025.12 출시)
        model = genai.GenerativeModel('gemini-3-flash-preview')
        
        prompt = f"""
        당신은 {topic_data['school']}의 전문 학교 소식지 에디터입니다.
        다음 정보를 바탕으로 생동감 있고 따뜻한 어조의 {topic_data['tone']} 기사를 작성해주세요.
        
        학년: {topic_data['grade']}
        행사명: {topic_data['event_name']}
        장소: {topic_data['location']}
        일시: {topic_data['date']}
        주요 키워드: {topic_data['keywords']}
        
        요구사항:
        1. 제목은 매력적으로 뽑아주세요. (첫 줄에 '제목: ' 형식으로)
        2. 본문은 400~600자 내외로 작성하세요.
        3. 문단은 보기 좋게 나누고 이모지를 적절히 사용하여 친근감을 주세요.
        4. 기사 끝에 관련 해시태그 5개를 작성해주세요. (형식: #태그1 #태그2 ...)
        """
        
        response = model.generate_content(prompt)
        text = response.text
        
        # 파싱
        title = topic_data['event_name']
        content = text
        hashtags = []
        
        # 해시태그 추출 logic
        found_hashtags = re.findall(r"#(\w+)", text)
        if found_hashtags:
            hashtags = found_hashtags[:5]
            # 텍스트에서 해시태그 부분 제거 (본문에는 깔끔하게만 남기기 위해)
            # 보통 마지막에 있으므로 마지막 줄 근처 처리
            lines = text.split('\n')
            clean_lines = [l for l in lines if not all(word.startswith('#') for word in l.split())]
            content = '\n'.join(clean_lines).strip()

        for line in text.split('\n'):
            if line.startswith("제목:") or line.startswith("##"):
                title = line.replace("제목:", "").replace("##", "").strip()
                # 이미 content를 위에서 hashtag 제거하며 세팅했으므로 다시 체크
                temp_content = content.replace(line, "").strip()
                if temp_content: content = temp_content
                break
                
        return title, content, hashtags
    except Exception as e:
        return f"AI 생성 오류", str(e), []

# ==========================================
# 3. 로직: 카드뉴스 생성 엔진 (Pillow)
# ==========================================
class CardNewsEngine:
    def __init__(self, theme_name):
        self.theme = THEMES[theme_name]
        self.size = (1080, 1080)
        self.font_path = FONT_PATH
        
    def _get_font(self, size, bold=False):
        # Note: NanumGothic-Regular doesn't have a separate bold file in this setup, 
        # but we use sized fonts for hierarchy.
        return ImageFont.truetype(self.font_path, size)

    def _draw_wrapped_text(self, draw, text, position, font, max_width, fill, max_lines=2):
        words = text.split()
        lines = []
        current_line = ""
        
        for word in words:
            test_line = current_line + (" " if current_line else "") + word
            bbox = draw.textbbox((0, 0), test_line, font=font)
            if bbox[2] - bbox[0] <= max_width:
                current_line = test_line
            else:
                lines.append(current_line)
                current_line = word
                if len(lines) >= max_lines: break
        
        if len(lines) < max_lines:
            lines.append(current_line)
        
        # If still too long, add ellipsis to the last line
        if len(lines) == max_lines:
            # Simple check if current_line was truncated or still too long
            last_line = lines[-1]
            bbox = draw.textbbox((0, 0), last_line, font=font)
            if bbox[2] - bbox[0] > max_width - 30:
                while last_line and draw.textbbox((0, 0), last_line + "...", font=font)[2] > max_width:
                    last_line = last_line[:-1]
                lines[-1] = last_line + "..."

        y = position[1]
        for line in lines:
            draw.text((position[0], int(y)), line, font=font, fill=fill)
            y += font.size * 1.4
        return int(y)

    def create_card(self, title, date, location, grade, hashtags, images):
        # 1. Canvas Setup (Clean White)
        canvas = Image.new("RGB", self.size, (255, 255, 255))
        draw = ImageDraw.Draw(canvas)
        
        main_rgb = self.theme["main"]
        accent_rgb = self.theme["accent"]
        
        # 2. Top Header (Profile Context)
        # 학년 배지 스타일
        badge_w = 120
        draw.rounded_rectangle((50, 50, 50 + badge_w, 100), radius=25, fill=main_rgb)
        draw.text((50 + 20, 60), grade if grade else "소식", font=self._get_font(22), fill=(255, 255, 255))
        
        # 학교/날짜 정보
        draw.text((50 + badge_w + 20, 55), "학교 소식지", font=self._get_font(24), fill=(30, 30, 30))
        draw.text((50 + badge_w + 20, 85), date, font=self._get_font(18), fill=(150, 150, 150))

        # 3. Main Title
        title_y = 140
        next_y = self._draw_wrapped_text(draw, title, (50, title_y), self._get_font(56), 980, (20, 20, 20), max_lines=2)
        
        # 4. Location Badge (이미지 바로 위에 배치)
        loc_y = next_y + 10
        if location:
            draw.text((50, loc_y), f"📍 {location}", font=self._get_font(28), fill=accent_rgb)
            img_y = loc_y + 60
        else:
            img_y = loc_y + 20

        # 5. Image Section
        img_w = 980 # 1000 -> 980 for better side margins
        img_h = 580 # 650 -> 580 to prevent bottom cutoff
        img_box = (50, int(img_y), 50 + img_w, int(img_y + img_h))
        
        # 박스 테두리 (Shadow 효과 느낌)
        draw.rectangle((48, int(img_y - 2), 52 + img_w, int(img_y + img_h + 2)), fill=(245, 245, 245))
        
        if images:
            self._render_image_grid(canvas, images, img_box)
            
        # 6. Hashtag Section
        # 이미지 바로 아래에 여백을 줄여서 배치
        tag_y = img_y + img_h + 30
        tag_str = " ".join([f"#{t}" for t in hashtags])
        self._draw_wrapped_text(draw, tag_str, (50, tag_y), self._get_font(34), 980, main_rgb, max_lines=2)
        
        # Footer Watermark (더 위로 올림)
        draw.text((820, 1010), "AI School Story", font=self._get_font(20), fill=(220, 220, 220))
        
        return canvas

    def _render_image_grid(self, canvas, image_paths, box):
        x, y, x2, y2 = box
        w, h = x2 - x, y2 - y
        count = len(image_paths)
        
        imgs = []
        for p in image_paths[:4]:
            try:
                # Handle both path strings and file-like objects
                if isinstance(p, str):
                    imgs.append(Image.open(p))
                else:
                    imgs.append(Image.open(io.BytesIO(p.getbuffer()) if hasattr(p, 'getbuffer') else p))
            except: continue
            
        if not imgs: return

        gap = 10
        if len(imgs) == 1:
            self._paste_cover(canvas, imgs[0], (x, y, w, h))
        elif len(imgs) == 2:
            half_w = (w - gap) // 2
            self._paste_cover(canvas, imgs[0], (x, y, half_w, h))
            self._paste_cover(canvas, imgs[1], (x + half_w + gap, y, half_w, h))
        elif len(imgs) == 3:
            big_w = (w * 2) // 3
            small_w = w - big_w - gap
            half_h = (h - gap) // 2
            self._paste_cover(canvas, imgs[0], (x, y, big_w, h))
            self._paste_cover(canvas, imgs[1], (x + big_w + gap, y, small_w, half_h))
            self._paste_cover(canvas, imgs[2], (x + big_w + gap, y + half_h + gap, small_w, half_h))
        else: # 4
            half_w = (w - gap) // 2
            half_h = (h - gap) // 2
            self._paste_cover(canvas, imgs[0], (x, y, half_w, half_h))
            self._paste_cover(canvas, imgs[1], (x + half_w + gap, y, half_w, half_h))
            self._paste_cover(canvas, imgs[2], (x, y + half_h + gap, half_w, half_h))
            self._paste_cover(canvas, imgs[3], (x + half_w + gap, y + half_h + gap, half_w, half_h))

    def _paste_cover(self, canvas, img, rect):
        rx, ry, rw, rh = rect
        iw, ih = img.size
        i_aspect = iw / ih
        r_aspect = rw / rh
        
        if i_aspect > r_aspect:
            new_h = int(rh)
            new_w = int(rh * i_aspect)
            img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
            left = (new_w - rw) // 2
            img = img.crop((int(left), 0, int(left + rw), int(rh)))
        else:
            new_w = int(rw)
            new_h = int(rw / i_aspect)
            img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
            top = (new_h - rh) // 2
            img = img.crop((0, int(top), int(rw), int(top + rh)))
            
        canvas.paste(img, (int(rx), int(ry)))

# ==========================================
# 3. 로직: PDF 생성 엔진
# ==========================================
class PDFEngine(FPDF):
    def __init__(self, theme_name, school_name):
        super().__init__()
        self.theme = THEMES[theme_name]
        self.school_name = school_name
        self.add_font("NanumGothic", "", FONT_PATH)
        # 상단 여백을 25mm로 설정하여 헤더(15mm)와 겹침 방지
        self.set_margins(10, 25, 10)
        self.set_auto_page_break(auto=True, margin=20)
        
    def header(self):
        # 표지(1페이지)가 아닐 때만 헤더 표시
        if self.page_no() > 1:
            try:
                # 배경 상자 (항상 페이지 최상단 0~15mm 위치)
                self.set_fill_color(*self.theme["main"])
                self.rect(0, 0, 210, 15, 'F')
                
                self.set_font("NanumGothic", "", 10)
                self.set_text_color(255, 255, 255)
                header_text = f"{self.school_name} 소식지"
                self.text(12, 10, header_text)
            except: pass

    def footer(self):
        if self.page_no() > 1:
            self.set_y(-15)
            self.set_font("NanumGothic", "", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f'- {self.page_no()} -', align='C')

    def draw_cover(self):
        self.add_page()
        # 표지는 상단 여백을 무시하고 전체 배경색 칠함
        self.set_fill_color(*self.theme["sub"])
        self.rect(0, 0, 210, 297, 'F')
        
        self.set_y(100)
        self.set_font("NanumGothic", "", 42)
        self.set_text_color(*self.theme["main"])
        self.cell(190, 30, self.school_name, align='C', ln=True)
        
        self.set_y(130)
        self.set_font("NanumGothic", "", 22)
        self.set_text_color(70, 70, 70)
        now = datetime.datetime.now()
        self.cell(190, 20, f"{now.year}학년도 {now.month}월 뉴스레터", align='C', ln=True)
        
        self.set_draw_color(*self.theme["accent"])
        self.set_line_width(1.5)
        self.line(60, 155, 150, 155)

    def calculate_article_height(self, article):
        """기사의 높이를 합산하여 반환"""
        h = 0
        self.set_font("NanumGothic", "", 16)
        title_lines = len(self.multi_cell(190, 10, str(article['title']), split_only=True))
        h += (title_lines * 10) + 10
        
        imgs_raw = article.get('images', '[]')
        imgs = json.loads(imgs_raw) if isinstance(imgs_raw, str) else imgs_raw
        if imgs:
            cnt = len(imgs)
            if cnt == 1: h += 95
            elif cnt <= 4: h += 140
        
        self.set_font("NanumGothic", "", 11)
        content_lines = len(self.multi_cell(190, 7, str(article['content']), split_only=True))
        h += (content_lines * 7) + 25
        return h

    def add_article(self, article):
        """본문 길이에 따라 사진 크기를 미세 조정하는 스마트 스케일링 적용"""
        h_no_img = self.calculate_article_height({**article, 'images': '[]'})
        imgs_raw = article.get('images', '[]')
        imgs = json.loads(imgs_raw) if isinstance(imgs_raw, str) else imgs_raw
        
        # 텍스트가 너무 길면 이미지를 좀 더 작게 조절 (기본 1.0 -> 0.7)
        scaling = 1.0
        if h_no_img > 150: # 본문이 페이지의 절반 이상을 차지하면
            scaling = 0.8
        
        h_final = self.calculate_article_height(article)
        if self.get_y() + h_final > 275:
            self.add_page()
        
        # 제목
        self.set_x(10)
        self.set_font("NanumGothic", "", 16)
        self.set_text_color(*self.theme["main"])
        self.multi_cell(190, 10, str(article.get('title', '')))
        self.ln(3)
        
        # 이미지 그리드 (스케일링 적용)
        if imgs:
            self.render_image_grid(imgs, scaling=scaling)
            self.ln(5)

        # 본문
        self.set_x(10)
        self.set_font("NanumGothic", "", 11)
        self.set_text_color(30, 30, 30)
        self.multi_cell(190, 7, str(article.get('content', '')))
        self.ln(15)

    def render_image_grid(self, imgs, scaling=1.0):
        cnt = len(imgs)
        try:
            if cnt == 1:
                w = 140 * scaling
                x = (210 - w) / 2
                self.image(imgs[0], x=x, w=w)
            elif cnt == 2:
                w = 92 * scaling
                curr_y = self.get_y()
                self.image(imgs[0], x=10, y=curr_y, w=w)
                self.image(imgs[1], x=108, y=curr_y, w=w)
                self.set_y(curr_y + (70 * scaling))
            elif cnt == 3:
                w_big = 110 * scaling
                w_small = 92 * scaling
                self.image(imgs[0], x=(210-w_big)/2, w=w_big)
                self.ln(5)
                curr_y = self.get_y()
                self.image(imgs[1], x=10, y=curr_y, w=w_small)
                self.image(imgs[2], x=108, y=curr_y, w=w_small)
                self.set_y(curr_y + (65 * scaling))
            else:
                w = 92 * scaling
                curr_y = self.get_y()
                self.image(imgs[0], x=10, y=curr_y, w=w)
                self.image(imgs[1], x=108, y=curr_y, w=w)
                self.ln(65 * scaling)
                new_y = self.get_y()
                self.image(imgs[2], x=10, y=new_y, w=w)
                self.image(imgs[3], x=108, y=new_y, w=w)
                self.set_y(new_y + (65 * scaling))
        except: pass

# ==========================================
# 4. 로직: Word 생성 엔진 (통합 본)
# ==========================================
def generate_docx_newsletter(articles, school_name, theme_name):
    from docx.shared import RGBColor
    theme = THEMES[theme_name]
    doc = Document()
    
    # 테마 색상 (제목용)
    main_color = RGBColor(*theme["main"])
    
    # 표지
    title = doc.add_heading(f'{school_name} 뉴스레터', 0)
    title.alignment = 1 
    
    now = datetime.datetime.now()
    p = doc.add_paragraph(f"{now.year}학년도 {now.month}월 소식")
    p.alignment = 1
    
    doc.add_page_break()

    for art in articles:
        # 기사 제목 (테마 색상 적용)
        h = doc.add_heading(str(art.get('title', '제목 없음')), level=1)
        for run in h.runs:
            run.font.color.rgb = main_color
        
        # 이미지 
        imgs_raw = art.get('images', '[]')
        imgs = json.loads(imgs_raw) if isinstance(imgs_raw, str) else imgs_raw
        if imgs:
            for img_p in imgs:
                try:
                    if os.path.exists(img_p):
                        doc.add_picture(img_p, width=doc.sections[0].page_width * 0.7)
                        doc.paragraphs[-1].alignment = 1
                except: pass
        
        # 본문
        doc.add_paragraph(str(art.get('content', '')))
        doc.add_page_break()
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 5. UI (Streamlit)
# ==========================================
def main():
    # 사이드바
    with st.sidebar:
        st.header("⚙️ 설정 (Settings)")
        api_key = st.text_input("Gemini API Key", type="password")
        if not api_key:
            st.warning("⚠️ API 키를 입력해야 기능을 사용할 수 있습니다.")
            with st.expander("🔑 API 키 발급 및 비용 안내", expanded=False):
                st.markdown("""
                이 시스템은 Google의 **Gemini 3 Flash** AI 모델을 사용하여 기사와 카드뉴스를 생성합니다.
                
                **1. 어디서 가져오나요?**
                - [Google AI Studio](https://aistudio.google.com/app/apikey)에서 누구나 구글 계정으로 즉시 발급 가능합니다.
                
                **2. 비용이 드나요?**
                - **아니오, 무료입니다.** 개인 개발 및 교육 목적의 무료 티어를 사용하면 별도의 결제 없이 무료로 이용할 수 있습니다.
                
                **3. 얼마나 쓸 수 있나요? (한도)**
                - **분당 15회 / 하루 1,500회** 호출이 가능합니다. 
                - 학교 소식지를 수십 번 다시 생성하더라도 **전혀 걱정 없이** 넉넉하게 사용할 수 있는 양입니다.
                
                **4. 발급 방법**
                1. [AI Studio](https://aistudio.google.com/app/apikey) 접속 ➔ 'Create API key' 클릭
                2. 발급된 키를 복사하여 위 칸에 넣어주세요.
                """)
        else:
            st.success("✅ API 키가 설정되었습니다.")
            
        school_name = st.text_input("학교명", value="서울디지털초등학교")
        
        # 테마 미리보기 적용
        st.write("🎨 디자인 테마 선택")
        theme_cols = st.columns([1, 4])
        selected_theme = st.selectbox("테마 목록", list(THEMES.keys()), label_visibility="collapsed")
        with theme_cols[0]:
            st.markdown(f'<div style="width:100%; height:30px; border-radius:5px; background:{THEMES[selected_theme]["hex"]}; border:1px solid #ddd;"></div>', unsafe_allow_html=True)
        with theme_cols[1]:
            st.caption(f"현재 테마: {selected_theme}")
        
        st.markdown("---")
        st.header("📂 기록 보관소")
        
        try:
            df_archive = pd.read_csv(DATA_FILE, encoding='utf-8-sig')
        except:
            df_archive = pd.DataFrame()

        st.write(f"총 {len(df_archive)}건의 기록이 있습니다.")
        
        tab_search, tab_batch = st.tabs(["🔍 개별 조회", "📰 뉴스레터 제작"])
        
        with tab_search:
            if not df_archive.empty:
                selection = st.selectbox("과거 기록 조회", 
                                       options=df_archive.index,
                                       format_func=lambda x: f"[{df_archive.iloc[x]['date']}] {df_archive.iloc[x]['event_name']}")
                if st.button("기록 보기", key="view_btn"):
                    st.session_state.current_view = df_archive.iloc[selection].to_dict()
            else:
                st.info("기록이 없습니다.")
        
        with tab_batch:
            if not df_archive.empty:
                selected_indices = st.multiselect("뉴스레터로 만들 기사 선택", 
                                            options=df_archive.index,
                                            format_func=lambda x: f"{df_archive.iloc[x]['event_name']}")
                
                if st.button("🚀 뉴스레터(PDF+Word) 통합 생성", use_container_width=True, type="primary"):
                    if selected_indices:
                        with st.spinner("두 가지 형식의 문서를 준비 중입니다..."):
                            articles = df_archive.iloc[selected_indices].to_dict('records')
                            
                            # 1. PDF 생성
                            pdf = PDFEngine(selected_theme, school_name)
                            pdf.draw_cover()
                            pdf.add_page()
                            for art in articles:
                                pdf.add_article(art)
                            pdf_data = bytes(pdf.output(dest='S'))
                            
                            # 2. Word 생성
                            docx_buffer = generate_docx_newsletter(articles, school_name, selected_theme)
                            
                            # 성공 후 세션에 저장하여 표시
                            st.session_state.ready_pdf = pdf_data
                            st.session_state.ready_docx = docx_buffer
                            st.success("문서 생성이 완료되었습니다! 아래에서 다운로드하세요.")

                # 생성된 파일이 있으면 다운로드 버튼 표시
                if 'ready_pdf' in st.session_state:
                    col_dl1, col_dl2 = st.columns(2)
                    with col_dl1:
                        st.download_button("📥 PDF 다운로드", 
                                         data=st.session_state.ready_pdf, 
                                         file_name="newsletter.pdf", 
                                         mime="application/pdf", 
                                         use_container_width=True)
                    with col_dl2:
                        st.download_button("📥 Word 다운로드", 
                                         data=st.session_state.ready_docx, 
                                         file_name="newsletter.docx", 
                                         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                                         use_container_width=True)
                    if st.button("🔄 새로 만들기", use_container_width=True):
                        del st.session_state.ready_pdf
                        del st.session_state.ready_docx
                        st.rerun()
            else:
                st.info("선택할 기록이 없습니다.")

        st.markdown("---")
        if not df_archive.empty:
            csv = df_archive.to_csv(index=False, encoding='utf-8-sig').encode('utf-8-sig')
            st.download_button("💾 전체 CSV 데이터 다운로드", data=csv, file_name="article_archive.csv", mime="text/csv")

    apply_custom_style()
    
    # 프로세스 스텝 바
    current_step = 0
    if 'draft_title' in st.session_state: current_step = 2
    elif 'current_view' in st.session_state: current_step = 3
    elif 'ready_pdf' in st.session_state: current_step = 3
    
    show_stepper(current_step)

    st.title("🎨 아이보리(AI-Story)")
    st.markdown("##### 📰 AI 기반 학교 소식지 제작 시스템")
    st.markdown("---")

    col1, col2 = st.columns([1.5, 1])

    with col1:
        st.subheader("📝 행사 정보 입력")
        with st.container(border=True):
            grade = st.selectbox("📌 참여 학년", ["전교생", "1학년", "2학년", "3학년", "4학년", "5학년", "6학년", "병설유치원", "기타"])
            event_name = st.text_input("📍 행사명", placeholder="예: 찾아가는 목공 체험")
            location = st.text_input("🏫 장소", placeholder="예: 학교 강당")
            event_date = st.date_input("📅 일시", datetime.date.today())
            tone = st.selectbox("🌈 기사 톤(분위기)", ["따뜻하고 감성적인", "활발하고 생동감 있는", "격조 있고 정중한", "간결하고 명확한"])
            keywords = st.text_area("✍️ 주요 활동 내용 (키워드)", placeholder="아이들이 나무 냄새를 맡으며 즐거워함, 뚝딱뚝딱 망치질 소리...", height=150)

        st.subheader("📸 사진 관리")
        with st.container(border=True):
            img_files = st.file_uploader("행사 사진들을 올려주세요 (여러 장 가능)", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True)
            
            selected_imgs = []
            if img_files:
                st.write("✅ 기사에 포함할 사진을 선택하세요")
                cols = st.columns(5) # 5개 컬럼으로 더 작게 표시
                for idx, img in enumerate(img_files):
                    with cols[idx % 5]:
                        st.image(img, use_container_width=True)
                        if st.checkbox(f"사진 {idx+1}", value=True, key=f"img_sel_{idx}"):
                            selected_imgs.append(img)
            else:
                st.info("사진을 올리면 기사와 함께 미리보기가 제공됩니다.")

    if st.button("✨ 기사 초안 생성하기", use_container_width=True, type="primary"):
        if not api_key:
            st.error("API 키가 없습니다. 사이드바에 키를 입력해주세요.")
        elif not event_name or not keywords:
            st.warning("행사명과 주요 내용을 입력해주세요.")
        else:
            with st.spinner("Gemini 3 Flash가 최신 기사를 작성 중입니다..."):
                input_data = {
                    "school": school_name,
                    "grade": grade,
                    "event_name": event_name,
                    "location": location,
                    "date": event_date.strftime("%Y/%m/%d"),
                    "tone": tone,
                    "keywords": keywords
                }
                title, content, hashtags = generate_article_gemini(api_key, input_data)
                
                st.session_state.draft_title = title
                st.session_state.draft_content = content
                st.session_state.draft_hashtags = hashtags
                st.session_state.current_input_data = input_data
                st.session_state.draft_images = selected_imgs
                st.success("AI 초안이 생성되었습니다! 아래에서 수정 후 카드뉴스를 만들어보세요.")

    # 편집 섹션 (초안이 있을 때만 표시)
    if 'draft_title' in st.session_state:
        st.markdown("---")
        st.subheader("💡 AI 생성 기사 수정하기")
        with st.container(border=True):
            edited_title = st.text_input("📝 기사 제목 수정", value=st.session_state.draft_title)
            edited_content = st.text_area("✍️ 기사 본문 수정", value=st.session_state.draft_content, height=300)
            
            # 사진 미리보기 (더 작고 컴팩트하게)
            if st.session_state.get('draft_images'):
                st.write("📸 업로드된 사진")
                imgs = st.session_state.draft_images
                cols = st.columns(5) # 5개 컬럼 고정으로 사진 크기 축소
                for idx, img in enumerate(imgs[:5]):
                    cols[idx % 5].image(img, use_container_width=True)
            
            edited_tags = st.text_input("🏷️ 해시태그 (공백으로 구분)", value=" ".join(st.session_state.get('draft_hashtags', [])))

            col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 1])
            with col_btn1:
                if st.button("📥 최종 기사 보관하기", use_container_width=True, type="primary"):
                    final_data = st.session_state.current_input_data
                    final_data.update({
                        "title": edited_title, 
                        "content": edited_content,
                        "hashtags": [t.strip() for t in edited_tags.split() if t.strip()]
                    })
                    save_to_archive(final_data, st.session_state.get('draft_images', []))
                    st.toast("저장되었습니다!")
                    del st.session_state.draft_title
                    st.rerun()
            
            with col_btn2:
                if st.button("🖼️ 카드뉴스 제작", use_container_width=True):
                    with st.spinner("요약 카드 생성 중..."):
                        engine = CardNewsEngine(selected_theme)
                        tags = [t.strip() for t in edited_tags.split() if t.strip()]
                        
                        temp_imgs = []
                        for img_file in st.session_state.get('draft_images', []):
                            tmp_path = f"tmp_{uuid.uuid4()}.png"
                            with open(tmp_path, "wb") as f:
                                f.write(img_file.getbuffer())
                            temp_imgs.append(tmp_path)
                        
                        # date 객체 또는 문자열 대응
                        input_info = st.session_state.current_input_data
                        raw_date = input_info['date']
                        if hasattr(raw_date, 'strftime'):
                            date_str = raw_date.strftime("%Y-%m-%d")
                        else:
                            date_str = str(raw_date)
                            
                        # 모든 메타데이터 포함하여 생성
                        summary_card = engine.create_card(
                            title=edited_title, 
                            date=date_str, 
                            location=input_info.get('location', ''),
                            grade=input_info.get('grade', ''),
                            hashtags=tags, 
                            images=temp_imgs
                        )
                            
                        for p in temp_imgs:
                            if os.path.exists(p): os.remove(p)
                            
                        st.session_state.preview_cards = [summary_card]
                        st.session_state.p_idx = 0
            
            with col_btn3:
                if st.button("❌ 취소 및 삭제", use_container_width=True):
                    del st.session_state.draft_title
                    st.rerun()

        # 카드뉴스 미리보기 UI (컬럼으로 크기 제한)
        if 'preview_cards' in st.session_state:
            st.markdown("---")
            st.subheader("🖼️ 카드뉴스 요약 카드")
            
            # 중앙 배치를 위해 3개 컬럼 사용 (좌우 여백 확보)
            _, col_mid, _ = st.columns([0.5, 1, 0.5])
            with col_mid:
                p_cards = st.session_state.preview_cards
                # Single Card Display
                st.image(p_cards[0], use_container_width=True)
                
                c_dl, c_close = st.columns(2)
                with c_dl:
                    img_io = io.BytesIO()
                    p_cards[0].save(img_io, format='PNG')
                    st.download_button("📥 PNG 다운로드", img_io.getvalue(), "card_news.png", "image/png", use_container_width=True)
                with c_close:
                    if st.button("✖️ 미리보기 닫기", use_container_width=True, key="close_preview_edit"):
                        del st.session_state.preview_cards
                        st.rerun()

    # 결과물 표시 및 개별 조회 화면
    if 'current_view' in st.session_state:
        st.markdown("---")
        v = st.session_state.current_view
        is_editing = st.session_state.get('is_editing', False)
        
        if is_editing:
            st.subheader(f"📝 기사 수정하기: {v.get('event_name', '')}")
            with st.container(border=True):
                new_title = st.text_input("📝 제목 수정", value=v['title'])
                st.write(f"**일시:** {v['date']} | **장소:** {v['location']} | **대상:** {v['grade']}")
                new_content = st.text_area("✍️ 본문 수정", value=v['content'], height=400)
                
                c1, c2 = st.columns(2)
                if c1.button("💾 변경사항 저장", use_container_width=True, type="primary"):
                    if update_archive(v['id'], new_title, new_content):
                        st.session_state.current_view['title'] = new_title
                        st.session_state.current_view['content'] = new_content
                        st.session_state.is_editing = False
                        st.success("수정사항이 저장되었습니다.")
                        st.rerun()
                if c2.button("🔙 취소", use_container_width=True):
                    st.session_state.is_editing = False
                    st.rerun()
        else:
            # 전문 매거진 스타일 레이아웃 (HTML/CSS 사용)
            import base64
            import textwrap

            def get_base64_img(path):
                try:
                    with open(path, "rb") as f:
                        data = base64.b64encode(f.read()).decode()
                        return f"data:image/png;base64,{data}"
                except: return ""

            imgs_raw = v.get('images', '[]')
            imgs = json.loads(imgs_raw) if isinstance(imgs_raw, str) else imgs_raw
            
            # 메타데이터 배지 HTML 생성
            badges_html = textwrap.dedent(f"""
                <div class="badge-container">
                    <div class="badge">📅 {v['date']}</div>
                    <div class="badge">🏫 {v['location']}</div>
                    <div class="badge">🎓 {v['grade']}</div>
                </div>
            """).strip()
            
            # 이미지 그리드 HTML 생성
            img_html = ""
            if imgs:
                img_count = min(len(imgs), 4)
                grid_class = f"grid-{img_count}"
                img_items_html = ""
                for i in range(img_count):
                    b64 = get_base64_img(imgs[i])
                    if b64:
                        img_items_html += f'<img src="{b64}" class="img-item">'
                
                img_html = f'<div class="img-grid {grid_class}">{img_items_html}</div>'

            # 기사 전체 렌더링 (인덴트 제거하여 마크다운 코드블록 방지)
            article_card_html = textwrap.dedent(f"""
                <div class="article-card">
                    <div class="article-header">
                        <div class="article-title">{v['title']}</div>
                        {badges_html}
                    </div>
                    {img_html}
                    <div class="article-content">{v['content']}</div>
                </div>
            """).strip()
            
            st.markdown(article_card_html, unsafe_allow_html=True)
            st.markdown("<br>", unsafe_allow_html=True)
            
            # --- 카드뉴스 및 공유/다운로드 영역 ---
            with st.container(border=True):
                st.markdown("### 🛠️ 추가 도구")
                col_tool1, col_tool2, col_tool3 = st.columns(3)
                
                with col_tool1:
                    if st.button("🖼️ 이 기사로 카드뉴스 만들기", use_container_width=True, type="primary"):
                        with st.spinner("요약 카드 생성 중..."):
                            engine = CardNewsEngine(selected_theme)
                            # 안전한 해시태그 파싱
                            tags_raw = v.get('hashtags', '[]')
                            try:
                                tags = json.loads(tags_raw) if isinstance(tags_raw, str) else tags_raw
                            except:
                                tags = []
                            
                            if not isinstance(tags, list): tags = []
                            
                            # 모든 메타데이터 포함하여 생성
                            summary_card = engine.create_card(
                                title=v['title'], 
                                date=v['date'], 
                                location=v.get('location', ''),
                                grade=v.get('grade', ''),
                                hashtags=tags, 
                                images=imgs
                            )
                            
                            st.session_state.preview_cards = [summary_card]
                            st.session_state.p_idx = 0
                
                with col_tool2:
                    # 공유용 텍스트 (안전하게 해시태그 처리)
                    tags_raw_v = v.get('hashtags', '[]')
                    try:
                        tags_v = json.loads(tags_raw_v) if isinstance(tags_raw_v, str) else tags_raw_v
                    except:
                        tags_v = []
                    if not isinstance(tags_v, list): tags_v = []
                    
                    tag_line = ' '.join([f'#{t}' for t in tags_v])
                    share_text = f"[{v['school']}] {v['title']}\n📅 일시: {v['date']}\n🏫 장소: {v['location']}\n\n{v['content'][:150]}...\n\n{tag_line}"
                    st.text_area("🔗 공유용 텍스트 (복사 가능)", value=share_text, height=100)
                
                with col_tool3:
                    # 개별 기사 PDF 다운로드 (임시 지원)
                    pdf_single = PDFEngine(selected_theme, school_name)
                    pdf_single.add_page()
                    pdf_single.add_article(v)
                    pdf_bytes = bytes(pdf_single.output(dest='S'))
                    st.download_button("📥 기사 PDF 다운로드", pdf_bytes, f"article_{v['id'][:8]}.pdf", "application/pdf", use_container_width=True)

            # 카드뉴스 미리보기 UI (컬럼으로 크기 제한)
            if 'preview_cards' in st.session_state:
                st.markdown("---")
                st.subheader("🖼️ 생성된 카드뉴스")
                
                _, col_mid_v, _ = st.columns([0.5, 1, 0.5])
                with col_mid_v:
                    p_cards = st.session_state.preview_cards
                    st.image(p_cards[0], use_container_width=True)
                    
                    c_zip, c_cls = st.columns(2)
                    with c_zip:
                        img_io = io.BytesIO()
                        p_cards[0].save(img_io, format='PNG')
                        st.download_button("📥 PNG 다운로드", img_io.getvalue(), "card_news.png", "image/png", use_container_width=True)
                    with c_cls:
                        if st.button("✖️ 카드뉴스 닫기", use_container_width=True, key="close_preview_view"):
                            del st.session_state.preview_cards
                            st.rerun()

            st.markdown("<br>", unsafe_allow_html=True)
            col_ctrl1, col_ctrl2 = st.columns(2)
            if col_ctrl1.button("✏️ 기사 수정하기", use_container_width=True):
                st.session_state.is_editing = True
                st.rerun()
            if col_ctrl2.button("✖️ 조회창 닫기", use_container_width=True):
                del st.session_state.current_view
                if 'is_editing' in st.session_state: del st.session_state.is_editing
                if 'preview_cards' in st.session_state: del st.session_state.preview_cards
                st.rerun()

if __name__ == "__main__":
    main()
